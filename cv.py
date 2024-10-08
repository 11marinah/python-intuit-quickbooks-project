from docx import Document
from docx.shared import Inches


document = Document()

# Profile picture
document.add_picture('user-profile-icon-free-vector.jpg',width=Inches(2.0))

# Personal information
name = input('Enter your name: ')
phone_number = int(input('Enter your phone number: '))
email = input('Enter your email: ')
document.add_paragraph(name + ' | ' + str(phone_number) + ' | ' + email)


# About me
document.add_heading('About me')
document.add_paragraph(input('Tell me about yourself: '))

# Work experience
document.add_heading('Work experience')
company = input('Enter you company name: ')
from_date = input('From date: ')
to_date = input('To date: ')
p = document.add_paragraph()
p.add_run(company + ': ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n')
experience = input('Enter your experience details: ')
p.add_run(experience)

# More experience
while True:
    has_more_experiences = input('Do you have more experience? yes or no? ')
    if has_more_experiences.lower() == 'yes':
        company = input('Enter you company name: ')
        from_date = input('From date: ')
        to_date = input('To date: ')
        p = document.add_paragraph()
        p.add_run(company + ': ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n')
        experience = input('Enter your experience details: ')
        p.add_run(experience)

    else:
        break

# Skills
document.add_heading('Skills')
skill = input('Enter your skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

# More skills
while True:
    has_more_skills = input('Do you have more skills? Yes or No? ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter your skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'    

    else:
        break


# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated by ioitms '


document.save('cv.docx')